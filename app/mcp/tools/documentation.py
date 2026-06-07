"""Documentation generation tools

Created by Sameer
"""
import os
import json
import logging
from datetime import datetime
from typing import Optional, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.mcp.server import register_tool
from app.services.salesforce import get_salesforce_connection

logger = logging.getLogger(__name__)

# Documents save path
DOCS_SAVE_PATH = os.getenv(
    "SFMCP_DOCS_PATH",
    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "documents")
)


# ── Styling helpers ────────────────────────────────────────────────────────────

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x58, 0x58, 0x58)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8EDF3"
HEADER_BG = "1B2A4A"


def _style_heading(paragraph, color=NAVY, size=14):
    """Apply consistent styling to a heading."""
    for run in paragraph.runs:
        run.font.color.rgb = color
        run.font.size = Pt(size)
        run.font.name = "Arial"


def _add_styled_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    sizes = {0: 22, 1: 16, 2: 13}
    colors = {0: NAVY, 1: NAVY, 2: BLUE}
    _style_heading(h, color=colors.get(level, NAVY), size=sizes.get(level, 14))
    return h


def _add_body_text(doc, text):
    """Add a styled body paragraph."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    return p


def _add_bullet(doc, text):
    """Add a styled bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    return p


def _add_styled_table(doc, headers, rows):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): HEADER_BG,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = LIGHT_BG if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg,
                qn('w:val'): 'clear'
            })
            shading.append(shading_elm)

    return table


def _add_divider(doc):
    """Add a visual divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("_" * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(6)


def _add_cover_page(doc, title, subtitle=""):
    """Add a professional cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = s.add_run(subtitle)
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = BLUE

    # Divider
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = d.add_run("━" * 40)
    run.font.color.rgb = BLUE
    run.font.size = Pt(12)

    # Date
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dt.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    # Page break
    doc.add_page_break()


def _ensure_save_path():
    """Ensure the documents directory exists."""
    os.makedirs(DOCS_SAVE_PATH, exist_ok=True)


def _make_file_path(prefix, file_name=None):
    """Generate a file path with auto-naming."""
    if not file_name:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"{prefix}_{timestamp}.docx"
    elif not file_name.endswith(".docx"):
        file_name += ".docx"
    return os.path.join(DOCS_SAVE_PATH, file_name), file_name


# ── BRD Tool ───────────────────────────────────────────────────────────────────

@register_tool
def generate_brd_document(
    summary: str,
    description: str,
    fields: Optional[List[str]] = None,
    triggers: Optional[List[str]] = None,
    flows: Optional[List[str]] = None,
    validations: Optional[List[str]] = None,
    acceptance_criteria: Optional[List[str]] = None,
    business_value: str = "",
    file_name: Optional[str] = None
) -> str:
    """Generate a Business Requirements Document (BRD) as a Word file (.docx). Parse the user requirement and pass each section as plain text strings. DO NOT pass code or scripts. The tool handles all formatting and styling internally.

    Args:
        summary: Short one-line summary of the requirement e.g. "Create a custom object for managing training records"
        description: Detailed plain text description of what needs to be built
        fields: List of field definitions as plain text e.g. ["Training Code (Auto Number)", "Mode (Picklist: Online, Classroom, Hybrid)"]
        triggers: List of trigger requirements as plain text e.g. ["When Training is inserted, if Duration is blank set to 1"]
        flows: List of flow requirements as plain text e.g. ["Send email notification when Training is created"]
        validations: List of validation rule requirements as plain text e.g. ["Duration must be greater than 0"]
        acceptance_criteria: List of acceptance criteria as plain text e.g. ["Training__c object must be created"]
        business_value: Plain text business justification for this requirement
        file_name: Custom file name (auto-generated if not provided)

    Returns:
        JSON with file path and document metadata
    """
    try:
        _ensure_save_path()
        file_path, file_name = _make_file_path("BRD", file_name)

        doc = Document()
        _add_cover_page(doc, "Business Requirements Document", summary)

        # Summary
        _add_styled_heading(doc, "Summary", level=1)
        _add_body_text(doc, summary)
        _add_divider(doc)

        # Description
        _add_styled_heading(doc, "Description", level=1)
        _add_body_text(doc, description)
        _add_divider(doc)

        # Fields to be Created
        if fields:
            _add_styled_heading(doc, "Fields to be Created", level=1)
            for field in fields:
                _add_bullet(doc, field)
            _add_divider(doc)

        # Automations
        has_automations = triggers or flows or validations
        if has_automations:
            _add_styled_heading(doc, "Automations", level=1)

            if triggers:
                _add_styled_heading(doc, "Triggers", level=2)
                for trigger in triggers:
                    _add_bullet(doc, trigger)

            if flows:
                _add_styled_heading(doc, "Flows", level=2)
                for flow in flows:
                    _add_bullet(doc, flow)

            if validations:
                _add_styled_heading(doc, "Validation Rules", level=2)
                for rule in validations:
                    _add_bullet(doc, rule)

            _add_divider(doc)

        # Acceptance Criteria
        if acceptance_criteria:
            _add_styled_heading(doc, "Acceptance Criteria", level=1)
            for criteria in acceptance_criteria:
                _add_bullet(doc, criteria)
            _add_divider(doc)

        # Business Value
        if business_value:
            _add_styled_heading(doc, "Business Value", level=1)
            _add_body_text(doc, business_value)

        doc.save(file_path)

        return json.dumps({
            "success": True,
            "message": "BRD document generated successfully",
            "file_path": os.path.abspath(file_path),
            "file_name": file_name,
            "doc_type": "BRD",
            "sections_included": {
                "summary": True,
                "description": True,
                "fields": bool(fields),
                "triggers": bool(triggers),
                "flows": bool(flows),
                "validations": bool(validations),
                "acceptance_criteria": bool(acceptance_criteria),
                "business_value": bool(business_value)
            }
        }, indent=2)

    except Exception as e:
        logger.exception("generate_brd_document failed")
        return json.dumps({"success": False, "error": str(e)})


# ── Design Document Tool ───────────────────────────────────────────────────────

@register_tool
def generate_design_document(
    title: str,
    requirement: str,
    solution: str,
    changes: List[str],
    components: Optional[List[str]] = None,
    objects_affected: Optional[List[str]] = None,
    dependencies: Optional[List[str]] = None,
    risks: Optional[List[str]] = None,
    file_name: Optional[str] = None
) -> str:
    """Generate a Design Document as a Word file (.docx). Parse the user requirement and pass each section as plain text strings. DO NOT generate code or scripts. The tool handles all formatting and styling internally.

    Args:
        title: Short title for the design e.g. "Case Closure Resolution Flow"
        requirement: Plain text description of the requirement being addressed
        solution: Plain text description of the proposed solution approach
        changes: List of changes needed as plain text e.g. ["Create custom field Resolution_Summary__c on Case", "Create Screen Flow Case_Closure_Flow"]
        components: List of Salesforce components involved e.g. ["LWC: caseClosureModal", "Flow: Case_Closure_Resolution", "Apex Trigger: CaseTrigger"]
        objects_affected: List of Salesforce objects affected e.g. ["Case", "Task", "EmailMessage"]
        dependencies: List of dependencies or prerequisites as plain text e.g. ["Case object must have Status field", "Email template must exist"]
        risks: List of risks or concerns as plain text e.g. ["Bulk case closure may hit governor limits", "Screen flow may not work in mobile"]
        file_name: Custom file name (auto-generated if not provided)

    Returns:
        JSON with file path and document metadata
    """
    try:
        _ensure_save_path()
        file_path, file_name = _make_file_path("Design_Document", file_name)

        doc = Document()
        _add_cover_page(doc, "Design Document", title)

        # Requirement
        _add_styled_heading(doc, "Requirement", level=1)
        _add_body_text(doc, requirement)
        _add_divider(doc)

        # Proposed Solution
        _add_styled_heading(doc, "Solution", level=1)
        _add_body_text(doc, solution)
        _add_divider(doc)

        # Changes Required
        _add_styled_heading(doc, "Changes", level=1)
        for change in changes:
            _add_bullet(doc, change)
        _add_divider(doc)

        # Components
        if components:
            _add_styled_heading(doc, "Components", level=1)
            table_rows = []
            for comp in components:
                parts = comp.split(":", 1) if ":" in comp else ["-", comp]
                table_rows.append([parts[0].strip(), parts[1].strip() if len(parts) > 1 else comp])
            _add_styled_table(doc, ["Type", "Name"], table_rows)
            _add_divider(doc)

        # Objects Affected
        if objects_affected:
            _add_styled_heading(doc, "Objects Affected", level=1)
            for obj in objects_affected:
                _add_bullet(doc, obj)
            _add_divider(doc)

        # Dependencies
        if dependencies:
            _add_styled_heading(doc, "Dependencies", level=1)
            for dep in dependencies:
                _add_bullet(doc, dep)
            _add_divider(doc)

        # Risks
        if risks:
            _add_styled_heading(doc, "Risks & Considerations", level=1)
            for risk in risks:
                _add_bullet(doc, risk)

        doc.save(file_path)

        return json.dumps({
            "success": True,
            "message": f"Design document generated: {title}",
            "file_path": os.path.abspath(file_path),
            "file_name": file_name,
            "doc_type": "Design",
            "sections_included": {
                "requirement": True,
                "solution": True,
                "changes": True,
                "components": bool(components),
                "objects_affected": bool(objects_affected),
                "dependencies": bool(dependencies),
                "risks": bool(risks)
            }
        }, indent=2)

    except Exception as e:
        logger.exception("generate_design_document failed")
        return json.dumps({"success": False, "error": str(e)})


# ── Test Document Tool ─────────────────────────────────────────────────────────

@register_tool
def generate_test_document(
    title: str,
    description: str,
    test_cases: List[str],
    preconditions: Optional[List[str]] = None,
    test_data: Optional[List[str]] = None,
    negative_tests: Optional[List[str]] = None,
    bulk_tests: Optional[List[str]] = None,
    file_name: Optional[str] = None
) -> str:
    """Generate a Test Document as a Word file (.docx). Parse the user requirement and pass each section as plain text strings. DO NOT generate code or scripts. The tool handles all formatting and styling internally.

    Args:
        title: Short title for the test plan e.g. "Training__c Object - Test Cases"
        description: Plain text description of what is being tested
        test_cases: List of test case descriptions as plain text e.g. ["Verify Training record can be created with all fields", "Verify Duration defaults to 1 when left blank"]
        preconditions: List of preconditions as plain text e.g. ["Training__c object must be deployed", "User must have Create permission"]
        test_data: List of test data requirements as plain text e.g. ["Create 5 Training records with various Mode values", "Create 1 Training with blank Duration"]
        negative_tests: List of negative test cases as plain text e.g. ["Enter text in Duration field - should show error", "Leave Training Name blank - should show required error"]
        bulk_tests: List of bulk/performance test cases as plain text e.g. ["Insert 200 Training records via Data Loader", "Bulk update Duration to blank on 500 records"]
        file_name: Custom file name (auto-generated if not provided)

    Returns:
        JSON with file path and document metadata
    """
    try:
        _ensure_save_path()
        file_path, file_name = _make_file_path("Test_Document", file_name)

        doc = Document()
        _add_cover_page(doc, "Test Document", title)

        # Description
        _add_styled_heading(doc, "Test Objective", level=1)
        _add_body_text(doc, description)
        _add_divider(doc)

        # Preconditions
        if preconditions:
            _add_styled_heading(doc, "Preconditions", level=1)
            for pre in preconditions:
                _add_bullet(doc, pre)
            _add_divider(doc)

        # Test Data
        if test_data:
            _add_styled_heading(doc, "Test Data Requirements", level=1)
            for td in test_data:
                _add_bullet(doc, td)
            _add_divider(doc)

        # Test Cases (as numbered table)
        _add_styled_heading(doc, "Test Cases", level=1)
        tc_rows = []
        for idx, tc in enumerate(test_cases, 1):
            tc_rows.append([str(idx), tc, "Pass / Fail"])
        _add_styled_table(doc, ["#", "Test Case", "Result"], tc_rows)
        _add_divider(doc)

        # Negative Tests
        if negative_tests:
            _add_styled_heading(doc, "Negative Test Cases", level=1)
            neg_rows = []
            for idx, nt in enumerate(negative_tests, 1):
                neg_rows.append([str(idx), nt, "Pass / Fail"])
            _add_styled_table(doc, ["#", "Test Case", "Result"], neg_rows)
            _add_divider(doc)

        # Bulk Tests
        if bulk_tests:
            _add_styled_heading(doc, "Bulk & Performance Tests", level=1)
            bulk_rows = []
            for idx, bt in enumerate(bulk_tests, 1):
                bulk_rows.append([str(idx), bt, "Pass / Fail"])
            _add_styled_table(doc, ["#", "Test Case", "Result"], bulk_rows)

        doc.save(file_path)

        total_tests = len(test_cases) + len(negative_tests or []) + len(bulk_tests or [])

        return json.dumps({
            "success": True,
            "message": f"Test document generated: {title}",
            "file_path": os.path.abspath(file_path),
            "file_name": file_name,
            "doc_type": "Test",
            "total_test_cases": total_tests,
            "sections_included": {
                "description": True,
                "preconditions": bool(preconditions),
                "test_data": bool(test_data),
                "test_cases": True,
                "negative_tests": bool(negative_tests),
                "bulk_tests": bool(bulk_tests)
            }
        }, indent=2)

    except Exception as e:
        logger.exception("generate_test_document failed")
        return json.dumps({"success": False, "error": str(e)})


# ── SF Object Schema Doc Tool ─────────────────────────────────────────────────

@register_tool
def generate_sf_object_documentation(
    object_name: str,
    include_fields: bool = True,
    include_relationships: bool = True,
    include_record_types: bool = True
) -> str:
    """Generate a Word document (.docx) documenting a Salesforce object schema, fields, relationships and record types by fetching metadata directly from the connected org.

    Args:
        object_name: Salesforce object API name (e.g. Account, Custom_Object__c)
        include_fields: Include field details in the document
        include_relationships: Include relationship information
        include_record_types: Include record type information

    Returns:
        JSON with file path and document metadata
    """
    try:
        sf = get_salesforce_connection()
        describe = sf.__getattr__(object_name).describe()

        _ensure_save_path()
        file_path, file_name = _make_file_path(f"{object_name}_Schema")

        doc = Document()
        _add_cover_page(doc, "Schema Documentation", f"{describe['label']} ({object_name})")

        # Overview
        _add_styled_heading(doc, "Overview", level=1)
        overview_rows = [
            ["API Name", object_name],
            ["Label", describe["label"]],
            ["Plural Label", describe.get("labelPlural", "N/A")],
            ["Custom Object", "Yes" if describe.get("custom") else "No"],
            ["Queryable", "Yes" if describe.get("queryable") else "No"],
        ]
        _add_styled_table(doc, ["Property", "Value"], overview_rows)
        _add_divider(doc)

        # Fields
        if include_fields:
            fields = describe.get("fields", [])
            _add_styled_heading(doc, f"Fields ({len(fields)})", level=1)
            field_rows = []
            for f in fields:
                field_rows.append([
                    f["name"],
                    f["label"],
                    f["type"],
                    "Yes" if not f.get("nillable", True) else "No",
                    "Yes" if f.get("custom", False) else "No"
                ])
            _add_styled_table(doc, ["API Name", "Label", "Type", "Required", "Custom"], field_rows)
            _add_divider(doc)

        # Relationships
        if include_relationships:
            child_rels = describe.get("childRelationships", [])
            _add_styled_heading(doc, f"Relationships ({len(child_rels)})", level=1)
            if child_rels:
                rel_rows = []
                for rel in child_rels[:30]:
                    rel_rows.append([
                        rel.get("relationshipName") or "N/A",
                        rel.get("childSObject", "N/A"),
                        rel.get("field", "N/A")
                    ])
                _add_styled_table(doc, ["Relationship", "Child Object", "Field"], rel_rows)
            else:
                _add_body_text(doc, "No child relationships found.")
            _add_divider(doc)

        # Record Types
        if include_record_types:
            record_types = describe.get("recordTypeInfos", [])
            if record_types:
                _add_styled_heading(doc, "Record Types", level=1)
                rt_rows = []
                for rt in record_types:
                    rt_rows.append([
                        rt.get("name", "N/A"),
                        "Yes" if rt.get("active") else "No",
                        "Yes" if rt.get("defaultRecordTypeMapping") else "No"
                    ])
                _add_styled_table(doc, ["Name", "Active", "Default"], rt_rows)

        doc.save(file_path)

        return json.dumps({
            "success": True,
            "message": f"Schema documentation generated for {object_name}",
            "file_path": os.path.abspath(file_path),
            "file_name": file_name,
            "object_name": object_name,
            "field_count": len(describe.get("fields", [])),
            "relationship_count": len(describe.get("childRelationships", []))
        }, indent=2)

    except Exception as e:
        logger.exception("generate_sf_object_documentation failed")
        return json.dumps({"success": False, "error": str(e)})
